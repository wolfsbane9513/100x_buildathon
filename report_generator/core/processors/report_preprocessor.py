from typing import List
from pathlib import Path
from datetime import datetime
from fpdf import FPDF
from docx import Document
from report_generator.app.config import config
import logging

logger = logging.getLogger(__name__)


class ReportProcessor:
    def __init__(self):
        self.output_dir = config.REPORTS_DIR
    
    def generate_report(self,
                       content: str,
                       output_format: str,
                       visualizations: List[str] = None) -> str:
        """Generate report in specified format"""
        if output_format not in config.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {output_format}")
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = self.output_dir / f"report_{timestamp}.{output_format}"
        
        if output_format == "pdf":
            self._generate_pdf(output_path, content, visualizations)
        elif output_format == "docx":
            self._generate_docx(output_path, content, visualizations)
        elif output_format == "html":
            self._generate_html(output_path, content, visualizations)
            
        return str(output_path)
    
    def _generate_pdf(self, output_path: Path, content: str, visualizations: List[str]):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, content)
        
        if visualizations:
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.cell(0, 10, "Visualizations", ln=True)

            for viz_path in visualizations:
                try:
                    if Path(viz_path).exists():
                        pdf.add_page()
                        pdf.image(viz_path, x=10, y=10, w=190)
                except Exception as e:
                    logger.error(f"Error adding visualization {viz_path} to PDF: {str(e)}")

        pdf.output(str(output_path))
    
    def _generate_docx(self, output_path: Path, content: str, visualizations: List[str]):
        doc = Document()
        doc.add_heading('Generated Report', 0)
        doc.add_paragraph(content)
        
        if visualizations:
            for viz_path in visualizations:
                doc.add_picture(viz_path, width=6000000)
        
        doc.save(output_path)
    
    def _generate_html(self, output_path: Path, content: str, visualizations: List[str]):
        html_content = ["<html><body>", f"<div>{content}</div>"]
        
        if visualizations:
            for viz_path in visualizations:
                html_content.append(
                    f'<img src="{viz_path}" style="max-width:100%;">'
                )
        
        html_content.append("</body></html>")
        output_path.write_text("\n".join(html_content))
