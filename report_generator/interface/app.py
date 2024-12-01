import gradio as gr
from typing import Dict, Any
import logging
import os
import json
from datetime import datetime
from report_generator.core.models.manager import ModelManager
from report_generator.core.agent import ReportGeneratorAgent
from report_generator.app.config import config
from fpdf import FPDF

logger = logging.getLogger(__name__)

class ReportGeneratorInterface:
    def __init__(self, ollama_manager=None):
        self.model_manager = ModelManager()
        self.ollama_manager = ollama_manager

    def create_interface(self):
        """Create Gradio interface"""
        with gr.Blocks(title="Report Generator Agent") as app:
            gr.Markdown("# Intelligent Report Generator")

            with gr.Tabs():
                # Model Selection Tab
                with gr.Tab("Model Selection"):
                    provider = gr.Radio(
                        choices=self.model_manager.get_available_providers(),
                        label="Select Model Provider",
                        interactive=True
                    )

                    model_name = gr.Radio(
                        choices=[],
                        label="Select Model",
                        interactive=True,
                        value=None
                    )

                    api_key = gr.Textbox(
                        label="API Key (if required)",
                        type="password",
                        visible=False,
                        interactive=True
                    )

                    model_info = gr.JSON(
                        label="Model Information",
                        visible=False
                    )

                # Data Source Tab
                with gr.Tab("Data Source"):
                    source_type = gr.Radio(
                        choices=["File Upload", "Database", "Both"],
                        label="Select Data Source",
                        interactive=True
                    )

                    with gr.Group(visible=True) as file_group:
                        file_input = gr.File(
                            label="Upload Data Files",
                            file_types=[".csv", ".json", ".xlsx", ".xls"],
                            file_count="multiple",
                            interactive=True
                        )

                    with gr.Group(visible=False) as db_group:
                        db_type = gr.Radio(
                            choices=["MongoDB", "MySQL", "PostgreSQL"],
                            label="Database Type",
                            interactive=True
                        )

                        connection_info = gr.JSON(
                            label="Connection Details",
                            visible=False
                        )

                        test_connection = gr.Button("Test Connection")

                # Generate Report Tab
                with gr.Tab("Generate Report"):
                    query = gr.Textbox(
                        label="What would you like to analyze?",
                        placeholder="Describe what you'd like to learn from the data...",
                        lines=3,
                        interactive=True
                    )

                    format_type = gr.Radio(
                        choices=["PDF", "HTML", "DOCX"],
                        label="Report Format",
                        value="PDF",
                        interactive=True
                    )

                    include_viz = gr.Checkbox(
                        label="Include Visualizations",
                        value=True,
                        interactive=True
                    )

                    generate_btn = gr.Button("Generate Report", variant="primary")

                    with gr.Row():
                        output_file = gr.File(label="Generated Report")
                        status = gr.Textbox(label="Status")

            def update_model_list(provider_value):
                """Update model list when provider changes"""
                try:
                    logger.info(f"Provider selected: {provider_value}")
                    if provider_value == "local":
                        models = self.model_manager.get_available_models(provider_value)
                        if models:
                            return (
                                gr.update(choices=models, value=None),
                                gr.update(visible=False),
                                gr.update(visible=True)
                            )
                        else:
                            return (
                                gr.update(choices=[], value=None),
                                gr.update(visible=False),
                                gr.update(visible=False)
                            )
                    elif provider_value == "openai":
                        models = ["gpt-3.5-turbo", "gpt-4"]
                        return (
                            gr.update(choices=models, value=None),
                            gr.update(visible=True),
                            gr.update(visible=True)
                        )
                    return gr.update(choices=[], value=None), gr.update(visible=False), gr.update(visible=False)
                except Exception as e:
                    logger.error(f"Error updating model list: {str(e)}")
                    return gr.update(choices=[], value=None), gr.update(visible=False), gr.update(visible=False)

            def update_data_source(source_type):
                """Update visible components based on data source selection"""
                show_files = source_type in ["File Upload", "Both"]
                show_db = source_type in ["Database", "Both"]

                return {
                    file_group: gr.update(visible=show_files),
                    db_group: gr.update(visible=show_db),
                    db_type: gr.update(visible=show_db),
                    connection_info: gr.update(visible=show_db),
                    test_connection: gr.update(visible=show_db)
                }

            async def generate_report(
                provider_value, model_name, api_key, source_type, files, db_info, query_text, format_type, include_viz
            ):
                """Handle report generation"""
                try:
                    if not model_name:
                        return None, "Please select a model first"

                    if not query_text:
                        return None, "Please enter a query"

                    if not files and source_type in ["File Upload", "Both"]:
                        return None, "Please upload at least one file"

                    os.makedirs('reports', exist_ok=True)

                    # Initialize the model based on user selection
                    model = self.model_manager.get_model(provider_value, model_name, api_key)

                    # Create a report generation agent with the selected model
                    agent = ReportGeneratorAgent(model)

                    context = {
                        'files': [f.name for f in files] if files else [],
                        'database': db_info if db_info else {},
                        'source_type': source_type
                    }

                    # Generate the report using the agent
                    result = await agent.generate_report(
                        query=query_text,
                        context=context,
                        output_format=format_type.lower(),
                        include_viz=include_viz
                    )

                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

                    if format_type.lower() == 'pdf':
                        return await save_pdf_report(result, timestamp)
                    elif format_type.lower() == 'docx':
                        return await save_docx_report(result, timestamp)
                    else:
                        return await save_html_report(result, timestamp)

                except Exception as e:
                    logger.error(f"Error in report generation: {str(e)}")
                    return None, f"Error generating report: {str(e)}"

            async def save_pdf_report(result: Dict[str, Any], timestamp: str) -> tuple:
                """Save report as PDF with enhanced wrapping, font adjustments, and visualizations."""
                try:
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_auto_page_break(auto=True, margin=15)
                    pdf.set_font("Arial", "B", 16)
                    pdf.cell(0, 10, "Generated Report", ln=True, align='C')
                    pdf.ln(10)

                    pdf.set_font("Arial", size=12)
                    content_lines = result['content'].split('\n')

                    for line in content_lines:
                        try:
                            pdf.multi_cell(0, 10, line)
                        except RuntimeError:
                            pdf.set_font("Arial", size=10)
                            pdf.multi_cell(0, 8, line)

                    pdf.ln(5)

                    if result.get('analysis'):
                        pdf.add_page()
                        pdf.set_font("Arial", "B", 14)
                        pdf.cell(0, 10, "Data Analysis", ln=True)
                        pdf.ln(5)

                        pdf.set_font("Arial", size=10)
                        analysis_text = json.dumps(result['analysis'], indent=2)
                        for line in analysis_text.split('\n'):
                            try:
                                pdf.multi_cell(0, 8, line)
                            except RuntimeError:
                                pdf.set_font("Arial", size=8)
                                pdf.multi_cell(0, 6, line)

                    # Add visualizations if present
                    if 'visualizations' in result:
                        for viz_path in result['visualizations']:
                            pdf.add_page()
                            pdf.image(viz_path, x=10, y=20, w=180)

                    output_path = f"reports/report_{timestamp}.pdf"
                    pdf.output(output_path)
                    return output_path, "Report generated successfully!"
                except Exception as e:
                    logger.error(f"Error saving PDF: {str(e)}")
                    raise

            async def save_docx_report(result, timestamp):
                """Save report as DOCX with visualizations."""
                try:
                    from docx import Document
                    from docx.shared import Inches

                    doc = Document()
                    doc.add_heading('Generated Report', 0)
                    doc.add_paragraph(result['content'])

                    if result.get('analysis'):
                        doc.add_heading('Data Analysis', level=1)
                        doc.add_paragraph(json.dumps(result['analysis'], indent=2))

                    # Add visualizations if present
                    if 'visualizations' in result:
                        doc.add_heading('Visualizations', level=1)
                        for viz_path in result['visualizations']:
                            doc.add_picture(viz_path, width=Inches(6))

                    output_path = f"reports/report_{timestamp}.docx"
                    doc.save(output_path)
                    return output_path, "Report generated successfully!"
                except Exception as e:
                    logger.error(f"Error saving DOCX: {str(e)}")
                    raise

            async def save_html_report(result, timestamp):
                """Save report as HTML, including visualizations."""
                try:
                    output_path = f"reports/report_{timestamp}.html"

                    # Properly handle the content, analysis, and visualizations for HTML format
                    html_content = """
                    <!DOCTYPE html>
                    <html>
                    <head><title>Generated Report</title></head>
                    <body>
                    <h1>Generated Report</h1>
                    <p>{content}</p>
                    <h2>Data Analysis</h2>
                    <pre>{analysis}</pre>
                    <h2>Visualizations</h2>
                    """.format(
                        content=result['content'].replace('\n', '<br>'),
                        analysis=json.dumps(result.get('analysis', {}), indent=2)
                    )

                    # Add visualizations if present
                    if 'visualizations' in result:
                        for viz_path in result['visualizations']:
                            html_content += f'<img src="{viz_path}" alt="Visualization" style="max-width:100%;height:auto;"><br>'

                    html_content += "</body></html>"

                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(html_content)

                    return output_path, "Report generated successfully!"
                except Exception as e:
                    logger.error(f"Error saving HTML: {str(e)}")
                    raise

            provider.change(fn=update_model_list, inputs=provider, outputs=[model_name, api_key, model_info])
            source_type.change(fn=update_data_source, inputs=source_type, outputs=[file_group, db_group, db_type, connection_info, test_connection])
            generate_btn.click(fn=generate_report, inputs=[provider, model_name, api_key, source_type, file_input, connection_info, query, format_type, include_viz], outputs=[output_file, status])

            return app

    def launch(self):
        """Launch the interface"""
        app = self.create_interface()
        app.launch(server_name="0.0.0.0", server_port=7860, share=True, max_threads=40)
