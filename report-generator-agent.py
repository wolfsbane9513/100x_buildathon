import os
from typing import List, Dict, Any, Optional
from pathlib import Path
import pandas as pd
import plotly.express as px
from datetime import datetime
from docx import Document
from fpdf import FPDF
import json
import gradio as gr
from llama_index.core import (
    Settings,
    VectorStoreIndex,
    get_response_synthesizer,
    SimpleDirectoryReader,
    StorageContext,
    load_index_from_storage,
)
from llama_index.agent import ReActAgent
from llama_index.tools import QueryEngineTool, PythonTool
from llama_index.core.tools import FunctionTool
from llama_index.readers.file import PDFReader, CSVReader
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.openai import OpenAI
from llama_index.core.query_engine import RetrieverQueryEngine
from llama_index.core.postprocessor import SimilarityPostprocessor

class DataProcessor:
    @staticmethod
    def read_csv(file_path: str) -> pd.DataFrame:
        return pd.read_csv(file_path)
    
    @staticmethod
    def read_json(file_path: str) -> Dict:
        with open(file_path, 'r') as f:
            return json.load(f)
    
    @staticmethod
    def read_excel(file_path: str) -> pd.DataFrame:
        return pd.read_excel(file_path)

class Visualizer:
    @staticmethod
    def create_bar_chart(data: pd.DataFrame, x: str, y: str, title: str) -> str:
        fig = px.bar(data, x=x, y=y, title=title)
        output_path = f"temp/bar_chart_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
        fig.write_image(output_path)
        return output_path
    
    @staticmethod
    def create_line_chart(data: pd.DataFrame, x: str, y: str, title: str) -> str:
        fig = px.line(data, x=x, y=y, title=title)
        output_path = f"temp/line_chart_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
        fig.write_image(output_path)
        return output_path
    
    @staticmethod
    def create_pie_chart(data: pd.DataFrame, names: str, values: str, title: str) -> str:
        fig = px.pie(data, names=names, values=values, title=title)
        output_path = f"temp/pie_chart_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
        fig.write_image(output_path)
        return output_path

class ReportGenerator:
    def __init__(self):
        self.embeddings = HuggingFaceEmbedding(model_name='BAAI/bge-small-en-v1.5')
        Settings.embed_model = self.embeddings
        self.llm = OpenAI(model="gpt-4")
        Settings.llm = self.llm
        self.data_processor = DataProcessor()
        self.visualizer = Visualizer()
        
    def _create_tools(self, data: pd.DataFrame) -> List[FunctionTool]:
        tools = [
            FunctionTool.from_defaults(
                fn=self.visualizer.create_bar_chart,
                name="create_bar_chart",
                description="Creates a bar chart from the data"
            ),
            FunctionTool.from_defaults(
                fn=self.visualizer.create_line_chart,
                name="create_line_chart",
                description="Creates a line chart from the data"
            ),
            FunctionTool.from_defaults(
                fn=self.visualizer.create_pie_chart,
                name="create_pie_chart",
                description="Creates a pie chart from the data"
            ),
            PythonTool.from_defaults(
                fn=lambda q: data.query(q),
                name="query_data",
                description="Query the dataset using pandas query syntax"
            ),
        ]
        return tools

    def create_report(self, 
                     data_path: str,
                     query: str,
                     output_format: str = "pdf",
                     include_visualizations: bool = True) -> str:
        # Read data based on file type
        file_extension = Path(data_path).suffix.lower()
        if file_extension == '.csv':
            data = self.data_processor.read_csv(data_path)
        elif file_extension == '.json':
            data = pd.DataFrame(self.data_processor.read_json(data_path))
        elif file_extension in ['.xlsx', '.xls']:
            data = self.data_processor.read_excel(data_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

        # Create agent with tools
        tools = self._create_tools(data)
        agent = ReActAgent.from_tools(tools, llm=self.llm, verbose=True)
        
        # Generate report content using agent
        response = agent.chat(f"""
        Generate a report based on the following request: {query}
        Include data analysis and insights.
        {"Include appropriate visualizations." if include_visualizations else ""}
        """)
        
        # Generate report in specified format
        output_path = self._generate_formatted_report(
            response.response,
            output_format,
            agent.get_tool_outputs()
        )
        
        return output_path

    def _generate_formatted_report(self,
                                 content: str,
                                 output_format: str,
                                 visualizations: List[str]) -> str:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = f"reports/report_{timestamp}.{output_format}"
        
        if output_format == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            # Add content
            pdf.multi_cell(0, 10, content)
            
            # Add visualizations
            for viz_path in visualizations:
                if viz_path.endswith('.png'):
                    pdf.add_page()
                    pdf.image(viz_path, x=10, y=10, w=190)
            
            pdf.output(output_path)
            
        elif output_format == "docx":
            doc = Document()
            doc.add_heading('Generated Report', 0)
            
            # Add content
            doc.add_paragraph(content)
            
            # Add visualizations
            for viz_path in visualizations:
                if viz_path.endswith('.png'):
                    doc.add_picture(viz_path, width=6000000)  # ~6 inches
            
            doc.save(output_path)
            
        elif output_format == "html":
            with open(output_path, 'w') as f:
                f.write("<html><body>")
                f.write(f"<div>{content}</div>")
                
                # Add visualizations
                for viz_path in visualizations:
                    if viz_path.endswith('.png'):
                        f.write(f'<img src="{viz_path}" style="max-width:100%;">')
                
                f.write("</body></html>")
        
        return output_path

def create_interface():
    report_generator = ReportGenerator()
    
    def generate_report(file, query, output_format, include_viz):
        try:
            temp_path = file.name
            return report_generator.create_report(
                temp_path,
                query,
                output_format,
                include_viz
            )
        except Exception as e:
            return f"Error generating report: {str(e)}"
    
    with gr.Blocks(title="Report Generator Agent") as app:
        gr.Markdown("# Intelligent Report Generator")
        
        with gr.Row():
            with gr.Column():
                file_input = gr.File(
                    label="Upload Data File (CSV, JSON, Excel)",
                    file_types=[".csv", ".json", ".xlsx", ".xls"]
                )
                query_input = gr.Textbox(
                    label="What kind of report do you need?",
                    placeholder="E.g., Generate a sales report for Q3 with monthly trends"
                )
                format_input = gr.Dropdown(
                    choices=["pdf", "docx", "html"],
                    value="pdf",
                    label="Output Format"
                )
                viz_input = gr.Checkbox(
                    label="Include Visualizations",
                    value=True
                )
                
            with gr.Column():
                output = gr.File(label="Generated Report")
                
        submit_btn = gr.Button("Generate Report")
        submit_btn.click(
            fn=generate_report,
            inputs=[file_input, query_input, format_input, viz_input],
            outputs=output
        )
        
    return app

if __name__ == "__main__":
    # Create necessary directories
    os.makedirs("temp", exist_ok=True)
    os.makedirs("reports", exist_ok=True)
    
    # Launch the interface
    app = create_interface()
    app.launch(server_name="0.0.0.0", server_port=7860)
